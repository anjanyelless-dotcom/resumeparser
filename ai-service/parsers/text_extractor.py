import os
import re
import logging
import time
from pathlib import Path
from typing import Dict, Optional, Tuple
import unicodedata

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available. Will use pymupdf as primary.")

try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("pymupdf not available. PDF extraction will be limited.")

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("Tesseract OCR not available. Scanned PDF processing will be limited.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX text extraction will be limited.")

try:
    from parsers.aws_textract_extractor import AwsTextractExtractor
    _textract_extractor = AwsTextractExtractor()  # singleton, reads env vars at import time
    TEXTRACT_AVAILABLE = _textract_extractor.is_available()
    if TEXTRACT_AVAILABLE:
        logging.info("✅ AWS Textract available and configured")
    else:
        logging.info("ℹ️ AWS Textract not configured (missing credentials or boto3) — will use local OCR")
except Exception as _textract_import_err:
    _textract_extractor = None
    TEXTRACT_AVAILABLE = False
    logging.warning(f"Could not load AwsTextractExtractor: {_textract_import_err}")

# PaddleOCR is not available for macOS ARM64, using Tesseract instead
PADDLEOCR_AVAILABLE = False
logging.info("ℹ️  PaddleOCR not available for macOS ARM64. Using Tesseract OCR for images.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """
    A comprehensive text extraction class that supports PDF, DOCX, and TXT files.
    Includes OCR fallback for scanned PDFs and text cleaning capabilities.
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.webp'}
        
    def extract_from_pdf(self, file_path: str, force_ocr: bool = False) -> Dict[str, any]:
        """
        Extract text from PDF using multi-tier strategy:
          Tier 1: pdfplumber            — best for text-based PDFs
          Tier 2: AWS Textract (cloud)  — scanned PDFs, table-heavy layouts
          Tier 3: pymupdf               — fallback digital text
          Tier 4: Tesseract OCR (local) — last-resort local OCR

        Args:
            file_path: Path to the PDF file
            force_ocr: Force AWS Textract (or local Tesseract if Textract unavailable)

        Returns:
            Dictionary with text, method_used, char_count, quality_score
        """
        MIN_CHAR_THRESHOLD = 200

        # ── force_ocr: prefer Textract, then Tesseract ──────────────────────
        if force_ocr:
            if TEXTRACT_AVAILABLE:
                try:
                    logger.info(f"[FORCE-OCR] Using AWS Textract: {file_path}")
                    result = _textract_extractor.extract_from_pdf(file_path)
                    char_count = result['char_count']
                    text = result['text']
                    logger.info(f"✅ Textract (force) extracted {char_count} chars")
                    return {
                        'text': text,
                        'method_used': 'aws_textract',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split())),
                        'has_tables': result.get('has_tables', False),
                    }
                except Exception as exc:
                    logger.error(f"Textract (force) failed: {exc}, falling back to Tesseract")

            if TESSERACT_AVAILABLE:
                try:
                    logger.info(f"[FORCE-OCR] Manual OCR trigger enabled. Using Tesseract: {file_path}")
                    text, avg_conf, quality_ok = self._extract_from_pdf_ocr(file_path)
                    char_count = len(text.strip())
                    logger.info(f"✅ OCR extraction completed: {char_count} chars, conf: {avg_conf:.1f}%, quality: {quality_ok}")
                    return {
                        'text': text,
                        'method_used': 'ocr',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split())),
                        'ocr_confidence': avg_conf,
                        'ocr_quality_ok': quality_ok
                    }
                except Exception as e:
                    logger.error(f"OCR extraction failed: {e}")
                    raise
            else:
                raise ImportError("force_ocr=True but neither AWS Textract nor Tesseract OCR is available.")

        # ── Tier 1: pdfplumber (text-based PDFs) ────────────────────────────
        if PDFPLUMBER_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with pdfplumber: {file_path}")
                text = self._extract_with_pdfplumber(file_path)
                char_count = len(text.strip())

                if char_count >= MIN_CHAR_THRESHOLD and self._is_text_quality_good(text):
                    logger.info(f"✅ pdfplumber extraction successful: {char_count} chars")
                    return {
                        'text': text,
                        'method_used': 'pdfplumber',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split()))
                    }
                elif char_count >= MIN_CHAR_THRESHOLD:
                    logger.warning(f"⚠️ pdfplumber extracted {char_count} chars but quality is poor — trying Textract")
                else:
                    logger.warning(f"⚠️ pdfplumber extracted only {char_count} chars (< {MIN_CHAR_THRESHOLD}) — trying Textract")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, trying Textract")

        # ── Tier 2: AWS Textract (scanned / table PDFs) ─────────────────────
        if TEXTRACT_AVAILABLE:
            try:
                logger.info(f"[TEXTRACT] Attempting PDF extraction with AWS Textract: {file_path}")
                result = _textract_extractor.extract_from_pdf(file_path)
                char_count = result['char_count']
                text = result['text']

                if char_count >= MIN_CHAR_THRESHOLD:
                    logger.info(f"✅ AWS Textract extraction successful: {char_count} chars, tables={result.get('has_tables', False)}")
                    return {
                        'text': text,
                        'method_used': 'aws_textract',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split())),
                        'has_tables': result.get('has_tables', False),
                    }
                else:
                    logger.warning(f"⚠️ Textract extracted only {char_count} chars (< {MIN_CHAR_THRESHOLD}), trying pymupdf")
            except Exception as exc:
                logger.warning(f"AWS Textract extraction failed: {exc}, trying pymupdf")

        # ── Tier 3: pymupdf (digital text fallback) ──────────────────────────
        if PYMUPDF_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with pymupdf: {file_path}")
                text = self._extract_with_pymupdf(file_path)
                char_count = len(text.strip())

                if char_count >= MIN_CHAR_THRESHOLD:
                    logger.info(f"✅ pymupdf extraction successful: {char_count} chars")
                    return {
                        'text': text,
                        'method_used': 'pymupdf',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split()))
                    }
                else:
                    logger.warning(f"⚠️ pymupdf extracted only {char_count} chars (< {MIN_CHAR_THRESHOLD}), trying local OCR")
            except Exception as e:
                logger.warning(f"pymupdf extraction failed: {e}, trying OCR")

        # ── Tier 4: Tesseract OCR (last-resort local OCR) ────────────────────
        if TESSERACT_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with Tesseract OCR: {file_path}")
                text, avg_conf, quality_ok = self._extract_from_pdf_ocr(file_path)
                char_count = len(text.strip())
                logger.info(f"✅ OCR extraction completed: {char_count} chars, conf: {avg_conf:.1f}%, quality: {quality_ok}")
                return {
                    'text': text,
                    'method_used': 'ocr',
                    'char_count': char_count,
                    'quality_score': self._calculate_quality_score(text, len(text.split())),
                    'ocr_confidence': avg_conf,
                    'ocr_quality_ok': quality_ok
                }
            except Exception as e:
                logger.error(f"OCR extraction failed: {e}")

        # All methods failed
        raise Exception(f"All PDF extraction methods failed for {file_path}. Install pdfplumber, boto3 (AWS), or tesseract.")
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """
        Extract text using pdfplumber (best for text-based PDFs).
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is not available")
        
        import pdfplumber
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        text = '\n'.join(text_parts)
        return self.clean_text(text)
    
    def _extract_with_pymupdf(self, file_path: str) -> str:
        """
        Extract text using pymupdf.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf is not available")
        
        text_parts = []
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        
        doc.close()
        text = '\n'.join(text_parts)
        return self.clean_text(text)
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image using OpenCV and NumPy:
        1. Grayscale
        2. Denoising using Bilateral Filter
        3. Deskewing (find text angle and rotate back)
        4. Otsu Threshold Binarization
        """
        try:
            import numpy as np
            import cv2
            
            img_np = np.array(image)
            # Convert RGB to BGR first if color image
            if len(img_np.shape) == 3:
                gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_np.copy()

            # Denoise
            denoised = cv2.bilateralFilter(gray, 9, 75, 75)

            # Find skew angle
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
            pts = cv2.findNonZero(thresh)
            angle = 0.0
            if pts is not None:
                rect = cv2.minAreaRect(pts)
                (cx, cy), (w, h), angle = rect
                if angle < -45:
                    angle = 90 + angle
                elif angle > 45:
                    angle = angle - 90
                
                if w < h:
                    angle = angle + 90 if angle < 0 else angle - 90

            # Rotate if skew is significant (0.5 to 20 degrees)
            if 0.5 < abs(angle) < 20:
                (h_img, w_img) = gray.shape[:2]
                center = (w_img // 2, h_img // 2)
                M = cv2.getRotationMatrix2D(center, angle, 1.0)
                rotated = cv2.warpAffine(denoised, M, (w_img, h_img), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_CONSTANT, borderValue=255)
                logger.info("Deskewed image by %.2f degrees", angle)
                denoised = rotated

            # Binarization
            _, binarized = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            return Image.fromarray(binarized)
        except Exception as exc:
            logger.warning("Image preprocessing failed in AI service: %s. Using original image.", exc)
            return image

    def _is_text_quality_good(self, text: str) -> bool:
        if not text or len(text) < 200:
            return False

        # Too many merged long tokens
        long_tokens = [w for w in text.split() if len(w) > 30]
        if len(long_tokens) > 5:
            return False

        # Too many weird uppercase/lowercase merges
        if len(re.findall(r"[a-z][A-Z]", text)) > 20:
            return False

        # Too many artificial separators — raised threshold; OCR/table PDFs legitimately use pipes
        if text.count("|") > 30:
            return False

        # Too many abnormal long alpha runs
        if len(re.findall(r"[A-Za-z]{25,}", text)) > 5:
            return False

        return True

    def _extract_from_pdf_ocr(self, file_path: str) -> Tuple[str, float, bool]:
        """
        Extract text from PDF using OCR as fallback.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (text, average_confidence, quality_ok)
        """
        if not TESSERACT_AVAILABLE:
            raise ImportError("Tesseract OCR is required for PDF OCR")
        
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf is required for OCR (to convert PDF to images)")
        
        try:
            text = ""
            doc = fitz.open(file_path)
            confidences = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # Preprocess image
                img = self._preprocess_image_for_ocr(img)
                
                # Perform OCR
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
                
                # Get word confidence scores
                try:
                    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                    conf_list = [float(c) for c in data.get("conf", []) if c is not None and float(c) != -1]
                    if conf_list:
                        confidences.extend(conf_list)
                except Exception as exc:
                    logger.debug("Failed to calculate confidence for page: %s", exc)
            
            doc.close()
            cleaned_text = self.clean_text(text)
            
            avg_conf = sum(confidences) / len(confidences) if confidences else 85.0
            quality_ok = self._is_text_quality_good(cleaned_text)
            
            return cleaned_text, avg_conf, quality_ok
            
        except Exception as e:
            logger.error(f"Error during OCR extraction from PDF {file_path}: {str(e)}")
            raise
    
    def extract_from_docx_with_styles(self, file_path: str) -> Tuple[str, Dict[str, bool]]:
        """
        Extract text from DOCX file and identify headings based on paragraph styles.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (text, heading_map) where heading_map has:
            - key: line text
            - value: True if paragraph has Heading 1 or Heading 2 style
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX text extraction")
        
        try:
            doc = Document(file_path)
            full_text = []
            heading_map = {}
            
            # Extract from paragraphs with style information
            for para in doc.paragraphs:
                if para.text.strip():
                    raw_text = para.text.strip()
                    full_text.append(raw_text)
                    
                    # Check if paragraph style is a heading
                    style_name = para.style.name if para.style else ''
                    is_heading = False
                    
                    # Check if style contains "Heading" and is level 1 or 2
                    if 'Heading' in style_name:
                        if 'Heading 1' in style_name or 'Heading 2' in style_name:
                            is_heading = True
                    
                    heading_map[raw_text] = is_heading
            
            # Extract from ALL table cells (no style info for tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                text = para.text.strip()
                                full_text.append(text)
                                heading_map[text] = False  # Table cells are not headings
            
            # Extract from text boxes inside shapes (no style info)
            try:
                from docx.oxml.ns import qn
                for shape in doc.inline_shapes:
                    try:
                        tb = shape._inline.graphic.graphicData.find('.//' + qn('wps:txbx'))
                        if tb is not None:
                            for para in tb.findall('.//' + qn('w:p')):
                                texts = [node.text for node in para.findall('.//' + qn('w:t')) if node.text]
                                if texts:
                                    text = ''.join(texts)
                                    full_text.append(text)
                                    heading_map[text] = False  # Text boxes are not headings
                    except:
                        pass
            except Exception as e:
                logger.debug(f"Could not extract text from shapes: {e}")
            
            # Join and clean the text
            raw_joined_text = '\n'.join(full_text)
            cleaned_text = self.clean_text(raw_joined_text)
            
            # Note: heading_map is based on raw text, but section splitter primarily uses
            # regex patterns to detect sections, so we return empty dict since the mapping
            # would be incorrect after line merging in clean_text()
            # Section splitter will rely on regex patterns like "PROFESSIONAL EXPERIENCE"
            
            return cleaned_text, {}
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file including paragraphs, tables, and text boxes.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX text extraction")
        
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract from ALL table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text.strip())
            
            # Extract from text boxes inside shapes
            try:
                from docx.oxml.ns import qn
                for shape in doc.inline_shapes:
                    try:
                        tb = shape._inline.graphic.graphicData.find('.//' + qn('wps:txbx'))
                        if tb is not None:
                            for para in tb.findall('.//' + qn('w:p')):
                                texts = [node.text for node in para.findall('.//' + qn('w:t')) if node.text]
                                if texts:
                                    full_text.append(''.join(texts))
                    except:
                        pass
            except Exception as e:
                logger.debug(f"Could not extract text from shapes: {e}")
            
            return self.clean_text('\n'.join(full_text))
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file with encoding detection.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text as string
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                # Last resort: try with errors='ignore'
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                logger.warning(f"Used error-ignoring encoding for {file_path}")
        
        return self.clean_text(text)
    
    def extract_from_image(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from image using Tesseract OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Dictionary with text, method_used, char_count, quality_score
        """
        logger.info(f"🖼️  Extracting text from image: {file_path}")
        start_time = time.time()
        
        # Use Tesseract OCR
        if TESSERACT_AVAILABLE:
            try:
                return self._extract_with_tesseract(file_path, start_time)
            except Exception as e:
                logger.error(f"❌ Tesseract OCR failed: {e}")
                raise
        else:
            raise ImportError("Tesseract OCR is not available. Please install it with: pip install pytesseract")
    
    def _extract_with_tesseract(self, file_path: str, start_time: float) -> Dict[str, any]:
        """Extract text using Tesseract OCR."""
        from PIL import Image
        
        # Open image and convert to grayscale
        image = Image.open(file_path)
        image = image.convert('L')
        
        # Perform OCR
        text = pytesseract.image_to_string(image, config='--psm 6')
        processing_time = time.time() - start_time
        
        # Clean OCR text
        text = self.clean_ocr_text(text)
        
        char_count = len(text.strip())
        logger.info(f"✅ Tesseract OCR extraction completed: {char_count} chars, time: {processing_time:.2f}s")
        
        return {
            'text': text,
            'method_used': 'tesseract',
            'char_count': char_count,
            'quality_score': self._calculate_quality_score(text, len(text.split())),
            'ocr_confidence': None,  # Tesseract doesn't provide confidence
            'processing_time_ms': processing_time * 1000
        }
    
    def clean_ocr_text(self, text: str) -> str:
        """
        Clean OCR output to remove noise and improve quality.
        
        Args:
            text: Raw OCR text
            
        Returns:
            Cleaned text
        """
        lines = text.split('\n')
        cleaned_lines = []
        prev_line = ""
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Skip duplicate lines
            if line == prev_line:
                continue
            
            # Skip very short lines (likely noise)
            if len(line) < 2:
                continue
            
            # Remove random special characters (OCR noise)
            line = re.sub(r'[^\w\s\-\.\,\;\:\(\)\[\]\{\}\/\@\#\$\%\&\*\+\=\_\!\?]', ' ', line)
            
            # Normalize whitespace
            line = re.sub(r'\s+', ' ', line)
            
            cleaned_lines.append(line)
            prev_line = line
        
        return '\n'.join(cleaned_lines)
    
    def extract(self, file_path: str, force_ocr: bool = False) -> Dict:
        """
        Extract text from file with automatic type detection and quality assessment.
        
        Args:
            file_path: Path to the file
            force_ocr: Whether to force OCR for PDF files
            
        Returns:
            Dictionary containing extracted text, method, word count, and quality score
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        text = ""
        method = ""
        
        try:
            if file_extension == '.pdf':
                # PDF extraction returns dict with metadata
                result = self.extract_from_pdf(str(file_path), force_ocr=force_ocr)
                text = result['text']
                method = result['method_used']
                word_count = len(text.split())
                quality_score = result['quality_score']
                
                logger.info(f"Successfully extracted text from {file_path.name} using {method}")
                logger.info(f"Characters: {result['char_count']}, Words: {word_count}, Quality: {quality_score:.2f}")
                
                return {
                    'text': text,
                    'method': method,
                    'word_count': word_count,
                    'quality_score': quality_score,
                    'char_count': result['char_count'],
                    'ocr_confidence': result.get('ocr_confidence'),
                    'ocr_quality_ok': result.get('ocr_quality_ok')
                }
            
            elif file_extension == '.docx':
                text = self.extract_from_docx(str(file_path))
                method = "python-docx"
            elif file_extension == '.txt':
                text = self.extract_from_txt(str(file_path))
                method = "direct"
            elif file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                # Image extraction returns dict with metadata
                result = self.extract_from_image(str(file_path))
                text = result['text']
                method = result['method_used']
                word_count = len(text.split())
                quality_score = result['quality_score']
                
                logger.info(f"Successfully extracted text from {file_path.name} using {method}")
                logger.info(f"Characters: {result['char_count']}, Words: {word_count}, Quality: {quality_score:.2f}")
                
                return {
                    'text': text,
                    'method': method,
                    'word_count': word_count,
                    'quality_score': quality_score,
                    'char_count': result['char_count'],
                    'ocr_confidence': result.get('ocr_confidence'),
                    'processing_time_ms': result.get('processing_time_ms')
                }
            
            # Calculate metrics for non-PDF files
            word_count = len(text.split())
            char_count = len(text.strip())
            quality_score = self._calculate_quality_score(text, word_count)
            
            logger.info(f"Successfully extracted text from {file_path.name} using {method}")
            logger.info(f"Characters: {char_count}, Words: {word_count}, Quality: {quality_score:.2f}")
            
            return {
                'text': text,
                'method': method,
                'word_count': word_count,
                'quality_score': quality_score,
                'char_count': char_count
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise
    
    def extract_with_font_metadata(self, file_path: str) -> Tuple[str, Dict[str, Dict]]:
        """
        Extract text along with font metadata from files.
        For PDF files, extracts font size, bold status, x position, and vertical gap.
        For DOCX and TXT files, returns text with empty metadata dictionary.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (text, font_metadata_dict) where font_metadata_dict has:
            - key: line text
            - value: dict with 'font_size', 'is_bold', 'x_position', 'vertical_gap'
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # For PDF files, extract with font metadata
        if file_extension == '.pdf':
            return self._extract_pdf_with_font_metadata(str(file_path))
        
        # For DOCX files, extract text with heading style information
        elif file_extension == '.docx':
            text, heading_map = self.extract_from_docx_with_styles(str(file_path))
            
            # Convert heading_map to font_metadata format for compatibility
            # Lines marked as headings get a pseudo-metadata entry
            font_metadata = {}
            for line_text, is_heading in heading_map.items():
                if is_heading:
                    # Mark as heading with high font size and bold
                    font_metadata[line_text] = {
                        'font_size': 14.0,  # Simulated large font
                        'is_bold': True,     # Headings are typically bold
                        'x_position': 0.0,   # Not applicable for DOCX
                        'vertical_gap': 0.0  # Not applicable for DOCX
                    }
            
            return text, font_metadata
        
        # For TXT files, extract text only (no font metadata)
        elif file_extension == '.txt':
            text = self.extract_from_txt(str(file_path))
            return text, {}
        
        raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_pdf_with_font_metadata(self, file_path: str) -> Tuple[str, Dict[str, Dict]]:
        """
        Extract text and font metadata from PDF using PyMuPDF.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Tuple of (text, font_metadata_dict)
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf is required for font metadata extraction")
        
        doc = fitz.open(file_path)
        all_lines = []
        font_metadata = {}
        prev_y_position = None
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Get text blocks with detailed information
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text_parts = []
                        font_sizes = []
                        is_bold = False
                        x_position = None
                        y_position = None
                        
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                line_text_parts.append(text)
                                font_sizes.append(span.get("size", 0))
                                
                                # Check if font name contains "Bold"
                                font_name = span.get("font", "")
                                if "Bold" in font_name or "bold" in font_name:
                                    is_bold = True
                                
                                # Get x position of first word in line
                                if x_position is None:
                                    x_position = span.get("bbox", [0, 0, 0, 0])[0]
                                
                                # Get y position for vertical gap calculation
                                if y_position is None:
                                    y_position = span.get("bbox", [0, 0, 0, 0])[1]
                        
                        if line_text_parts:
                            line_text = " ".join(line_text_parts)
                            all_lines.append(line_text)
                            
                            # Calculate average font size
                            avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0
                            
                            # Calculate vertical gap from previous line
                            vertical_gap = 0
                            if prev_y_position is not None and y_position is not None:
                                vertical_gap = abs(y_position - prev_y_position)
                            
                            # Store metadata for this line
                            font_metadata[line_text] = {
                                'font_size': round(avg_font_size, 2),
                                'is_bold': is_bold,
                                'x_position': round(x_position, 2) if x_position is not None else 0,
                                'vertical_gap': round(vertical_gap, 2)
                            }
                            
                            prev_y_position = y_position
        
        doc.close()
        
        # Combine all lines into text
        text = '\n'.join(all_lines)
        text = self.clean_text(text)
        
        return text, font_metadata
    
    def calculate_baseline_font_size(self, font_metadata: Dict[str, Dict]) -> float:
        """
        Calculate the baseline (most common) font size from font metadata.
        This represents the body text size for the resume.
        
        Args:
            font_metadata: Dictionary with line text as keys and metadata dicts as values
            
        Returns:
            Most common font size, or 11.0 if no metadata exists
        """
        if not font_metadata:
            return 11.0
        
        # Collect all font sizes
        font_sizes = []
        for metadata in font_metadata.values():
            if 'font_size' in metadata and metadata['font_size'] > 0:
                font_sizes.append(metadata['font_size'])
        
        if not font_sizes:
            return 11.0
        
        # Find the most common font size (mode)
        from collections import Counter
        font_size_counts = Counter(font_sizes)
        most_common_size = font_size_counts.most_common(1)[0][0]
        
        return most_common_size
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing sensitive information and normalizing format.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # DO NOT remove emails/phones here - they need to be extracted first!
        # Privacy removal should happen after parsing, not before
        
        # Normalize whitespace - preserve newlines
        # First normalize multiple newlines to double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Then normalize spaces on same line (but keep newlines)
        lines = text.split('\n')
        lines = [re.sub(r'[ \t]+', ' ', line) for line in lines]
        text = '\n'.join(lines)
        
        # FIX 1: Merge lines that end with abbreviations (SR., JR., DR., etc.)
        # This prevents "SR.\nPOWER BI DEVELOPER" from being split (regardless of next line length)
        abbreviation_pattern = r'\b(SR|JR|DR|MR|MS|MRS|PROF|REV|HON|ESQ|PHD|MD|DDS|DVM)\.\s*\n\s*'
        text = re.sub(abbreviation_pattern, r'\1. ', text, flags=re.IGNORECASE)
        
        # FIX 2: Also merge if a line is ONLY an abbreviation (e.g., "SR.\n" on its own line)
        # Pattern: line with only abbreviation followed by newline
        standalone_abbrev_pattern = r'^(SR|JR|DR|MR|MS|MRS|PROF|REV|HON|ESQ|PHD|MD|DDS|DVM)\.\s*$'
        lines = text.split('\n')
        merged_abbrev_lines = []
        i = 0
        while i < len(lines):
            current_line = lines[i].strip()
            
            # Check if current line is ONLY an abbreviation
            if re.match(standalone_abbrev_pattern, current_line, re.IGNORECASE):
                # Merge with next line if it exists
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line:
                        merged_abbrev_lines.append(f"{current_line} {next_line}")
                        i += 2
                        continue
            
            merged_abbrev_lines.append(lines[i])
            i += 1
        
        text = '\n'.join(merged_abbrev_lines)
        
        # NOTE: Short-line merging was removed. AWS Textract and Tesseract OCR return
        # each detected text element on its own line (title, company, date are all short).
        # Merging short lines destroys the structure needed for section parsing.
        
        # Remove non-printable characters except newlines and tabs
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize unicode characters
        text = unicodedata.normalize('NFKC', text)
        
        # Remove excessive punctuation (keep @ for emails)
        text = re.sub(r'[^\w\s\n\t.,;:!?()[\]{}"\'\-@]', '', text)
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        text = re.sub(r'([.,;:!?])\s+', r'\1 ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _calculate_quality_score(self, text: str, word_count: int) -> float:
        """
        Calculate quality score based on text characteristics.
        
        Args:
            text: Extracted text
            word_count: Number of words in text
            
        Returns:
            Quality score between 0 and 1
        """
        if not text or word_count == 0:
            return 0.0
        
        score = 0.0
        
        # Base score from text length (0-0.4)
        length_score = min(word_count / 500, 1.0) * 0.4
        score += length_score
        
        # Readability score based on average word length (0-0.2)
        words = text.split()
        if words:
            avg_word_length = sum(len(word) for word in words) / len(words)
            readability_score = min(avg_word_length / 6, 1.0) * 0.2
            score += readability_score
        
        # Structure score based on paragraphs and newlines (0-0.2)
        paragraphs = text.split('\n\n')
        structure_score = min(len(paragraphs) / 10, 1.0) * 0.2
        score += structure_score
        
        # Content diversity score based on unique words (0-0.2)
        unique_words = set(word.lower() for word in words if len(word) > 3)
        if words:
            diversity_score = min(len(unique_words) / len(words), 1.0) * 0.2
            score += diversity_score
        
        return min(score, 1.0)
    
    def get_supported_formats(self) -> list:
        """
        Get list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return list(self.supported_extensions)
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        return Path(file_path).suffix.lower() in self.supported_extensions


# Import required for OCR image processing
import io

# Example usage and testing
if __name__ == "__main__":
    # Test the extractor
    extractor = TextExtractor()
    
    # Test with a sample file (if available)
    test_file = "sample.pdf"  # Replace with actual test file
    
    if os.path.exists(test_file):
        try:
            result = extractor.extract(test_file)
            print(f"Extraction successful!")
            print(f"Method: {result['method']}")
            print(f"Word count: {result['word_count']}")
            print(f"Quality score: {result['quality_score']:.2f}")
            print(f"Text preview: {result['text'][:200]}...")
        except Exception as e:
            print(f"Extraction failed: {e}")
    else:
        print(f"Test file {test_file} not found")
        print(f"Supported formats: {extractor.get_supported_formats()}")
